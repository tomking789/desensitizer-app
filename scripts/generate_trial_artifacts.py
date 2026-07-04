from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    samples = root / "demo" / "示例文件"
    marketing = root / "marketing"
    samples.mkdir(parents=True, exist_ok=True)
    marketing.mkdir(parents=True, exist_ok=True)

    (samples / "客户访谈记录.txt").write_text(
        "客户姓名: 张三\n"
        "手机号: 13812345678\n"
        "邮箱: zhangsan@example.com\n"
        "身份证号: 110101199003078899\n"
        "客户公司: 星河智造有限公司\n"
        "项目名称: 北区渠道优化项目\n"
        "合同编号: HT-2026-0703-001\n"
        "访谈摘要: 客户计划将内部销售数据交给 AI 做分析，要求先隐藏联系人、合同号和公司名称。\n",
        encoding="utf-8",
    )

    (samples / "客户清单.csv").write_text(
        "客户姓名,手机号,邮箱,公司,项目编号\n"
        "李四,13900001111,lisi@example.com,星河智造有限公司,PRJ-2026-001\n"
        "王五,13700002222,wangwu@example.com,北辰咨询有限公司,PRJ-2026-002\n",
        encoding="utf-8-sig",
    )

    document = Document()
    document.add_heading("项目合同摘要", level=1)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei UI"
            run.font.size = Pt(14)
    document.add_paragraph("甲方：星河智造有限公司")
    document.add_paragraph("联系人：张三，手机号：13812345678，邮箱：zhangsan@example.com")
    document.add_paragraph("合同编号：HT-2026-0703-001")
    document.add_paragraph("项目名称：北区渠道优化项目")
    document.add_paragraph("摘要：本文件用于演示在交给 AI 分析前，如何先对客户名称、联系人、手机号、邮箱和合同编号脱敏。")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "字段"
    header[1].text = "原始示例"
    header[2].text = "说明"
    for row in [
        ("客户", "星河智造有限公司", "企业名称"),
        ("联系人", "张三", "个人姓名"),
        ("手机号", "13812345678", "个人信息"),
    ]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.save(samples / "项目合同摘要.docx")

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "客户资料"
    worksheet.append(["客户姓名", "手机号", "邮箱", "身份证号", "公司", "合同编号"])
    worksheet.append(["张三", "13812345678", "zhangsan@example.com", "110101199003078899", "星河智造有限公司", "HT-2026-0703-001"])
    worksheet.append(["李四", "13900001111", "lisi@example.com", "110101198802023456", "北辰咨询有限公司", "HT-2026-0703-002"])
    for col in range(1, 7):
        worksheet.column_dimensions[chr(64 + col)].width = 24
    workbook.save(samples / "客户资料表.xlsx")

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CNTitle", fontName="STSong-Light", fontSize=22, leading=28, textColor=colors.HexColor("#17324d"), spaceAfter=10))
    styles.add(ParagraphStyle(name="CNBody", fontName="STSong-Light", fontSize=9.5, leading=14, textColor=colors.HexColor("#263238")))

    sample_pdf = samples / "合同摘要.pdf"
    doc = SimpleDocTemplate(str(sample_pdf), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    story = [
        Paragraph("合同摘要示例", styles["CNTitle"]),
        Paragraph("甲方：星河智造有限公司", styles["CNBody"]),
        Paragraph("联系人：张三，手机号：13812345678，邮箱：zhangsan@example.com", styles["CNBody"]),
        Paragraph("合同编号：HT-2026-0703-001", styles["CNBody"]),
        Paragraph("项目名称：北区渠道优化项目", styles["CNBody"]),
        Spacer(1, 8),
        Paragraph("该 PDF 是文字版 PDF，用于演示原版式覆盖式脱敏。扫描版 PDF 需要先自行转换为文字版。", styles["CNBody"]),
    ]
    doc.build(story)

    product_pdf = marketing / "本地资料脱敏工具-产品介绍.pdf"
    c = canvas.Canvas(str(product_pdf), pagesize=A4)
    width, height = A4
    c.setFillColor(colors.HexColor("#eef5fb"))
    c.rect(0, 0, width, height, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#17324d"))
    c.setFont("STSong-Light", 24)
    c.drawString(22 * mm, height - 28 * mm, "本地资料脱敏工具")
    c.setFont("STSong-Light", 12)
    c.setFillColor(colors.HexColor("#546e7a"))
    c.drawString(22 * mm, height - 37 * mm, "AI 办公前的本地安全前置工具 | 专业试用版")

    c.setFillColor(colors.white)
    c.roundRect(20 * mm, height - 95 * mm, 170 * mm, 42 * mm, 6, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#1769aa"))
    c.setFont("STSong-Light", 13)
    c.drawString(27 * mm, height - 66 * mm, "核心价值")
    c.setFillColor(colors.HexColor("#263238"))
    c.setFont("STSong-Light", 10)
    for i, line in enumerate(
        [
            "在把 Word、Excel、PDF 和文本资料交给 ChatGPT、Codex 或其他 AI 工具前，",
            "先在本机识别、确认、替换敏感信息，并生成可审计报告。",
            "默认加密映射表，降低原始信息外泄风险。",
        ]
    ):
        c.drawString(27 * mm, height - (76 + i * 7) * mm, line)

    cards = [
        ("本地处理", "默认在本机处理支持文件，不主动上传资料。"),
        ("批量脱敏", "支持多文件、文件夹和递归子文件夹。"),
        ("人工确认", "自动识别后可启用、禁用、修改和补充。"),
        ("可还原", "通过加密映射表按需还原脱敏文件。"),
    ]
    x0, y0 = 20 * mm, height - 155 * mm
    for idx, (title, body) in enumerate(cards):
        x = x0 + (idx % 2) * 87 * mm
        y = y0 - (idx // 2) * 38 * mm
        c.setFillColor(colors.white)
        c.roundRect(x, y, 82 * mm, 30 * mm, 5, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#1769aa"))
        c.setFont("STSong-Light", 12)
        c.drawString(x + 7 * mm, y + 19 * mm, title)
        c.setFillColor(colors.HexColor("#263238"))
        c.setFont("STSong-Light", 8.8)
        c.drawString(x + 7 * mm, y + 10 * mm, body)

    c.setFillColor(colors.white)
    c.roundRect(20 * mm, 45 * mm, 170 * mm, 44 * mm, 6, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#1769aa"))
    c.setFont("STSong-Light", 13)
    c.drawString(27 * mm, 76 * mm, "适用场景")
    c.setFillColor(colors.HexColor("#263238"))
    c.setFont("STSong-Light", 9.5)
    for i, scene in enumerate(["律师/法务材料 AI 摘要", "财税审计资料 AI 分析", "HR 简历和员工资料处理", "咨询访谈纪要整理", "企业内部文档交给 AI 前预处理"]):
        c.drawString(29 * mm, (67 - i * 6) * mm, "• " + scene)

    c.setFillColor(colors.HexColor("#546e7a"))
    c.setFont("STSong-Light", 8)
    c.drawString(22 * mm, 24 * mm, "提示：自动识别无法保证 100% 覆盖，高风险资料应人工复核；如需更新、模板和支持服务，请联系出品方。")
    c.setFillColor(colors.HexColor("#17324d"))
    c.drawRightString(188 * mm, 18 * mm, "艺林万象（北京）科技有限公司")
    c.save()

    print("generated trial artifacts")


if __name__ == "__main__":
    main()
